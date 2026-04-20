from __future__ import annotations

import io
from html.parser import HTMLParser

from fastapi import APIRouter
from fastapi.responses import Response
from pydantic import BaseModel, Field

router = APIRouter()


class ExportDocxPayload(BaseModel):
    title: str = Field(default="Document")
    html: str = Field(default="")


class _HtmlToLinesParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.lines: list[tuple[str, str]] = []
        self.current: list[str] = []
        self.block = "p"
        self.list_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in {"p", "div", "h1", "h2", "h3", "blockquote"}:
            self._flush()
            self.block = tag
        elif tag in {"ul", "ol"}:
            self.list_depth += 1
        elif tag == "li":
            self._flush()
            self.block = "li"
        elif tag == "br":
            self.current.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"p", "div", "h1", "h2", "h3", "blockquote", "li"}:
            self._flush()
            self.block = "p"
        elif tag in {"ul", "ol"} and self.list_depth > 0:
            self.list_depth -= 1

    def handle_data(self, data: str) -> None:
        text = data.replace("\xa0", " ")
        if text:
            self.current.append(text)

    def _flush(self) -> None:
        text = "".join(self.current).strip()
        self.current = []
        if text:
            self.lines.append((self.block, text))


@router.post("/document/export-docx")
async def export_docx(payload: ExportDocxPayload):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except Exception as exc:
        return Response(
            content=f"python-docx is not available: {exc}",
            media_type="text/plain",
            status_code=500,
        )

    parser = _HtmlToLinesParser()
    parser.feed(payload.html or "")
    document = Document()
    document.core_properties.title = payload.title
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    def apply_run_font(run, *, size: int, bold: bool = False, italic: bool = False, color: tuple[int, int, int] | None = None):
        run.bold = bold
        run.italic = italic
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if color is not None:
            run.font.color.rgb = RGBColor(*color)

    for block, text in parser.lines or [("p", payload.title)]:
        if block == "h1":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(10)
            run = paragraph.add_run(text)
            apply_run_font(run, size=22, bold=True, color=(17, 24, 39))
        elif block == "h2":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(text)
            apply_run_font(run, size=16, bold=True, color=(17, 24, 39))
        elif block == "h3":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(text)
            apply_run_font(run, size=13, bold=True, color=(17, 24, 39))
        elif block == "blockquote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(18)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(text)
            apply_run_font(run, size=11, italic=True, color=(71, 85, 105))
        elif block == "li":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(18)
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(f"• {text}")
            apply_run_font(run, size=11)
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(text)
            apply_run_font(run, size=11, color=(31, 41, 55))

    buffer = io.BytesIO()
    document.save(buffer)
    filename = f"{payload.title or 'document'}.docx"
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
