"""Knowledge base core: ingestion, embeddings, Qdrant storage, and search."""
from __future__ import annotations

import asyncio
from collections import Counter
import hashlib
import io
import logging
import math
import re
import uuid
from dataclasses import dataclass
from typing import TYPE_CHECKING

from qdrant_client import AsyncQdrantClient
from qdrant_client.models import (
    Distance,
    FieldCondition,
    Filter,
    Fusion,
    FusionQuery,
    MatchValue,
    Modifier,
    PointStruct,
    Prefetch,
    SparseVector,
    SparseVectorParams,
    VectorParams,
)
from sqlalchemy import delete, select, update

if TYPE_CHECKING:
    import httpx
    from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)

_manager: KnowledgeManager | None = None


def get_knowledge_manager() -> "KnowledgeManager | None":
    return _manager


async def shutdown_knowledge_manager() -> None:
    """Stop background workers for the active knowledge manager."""
    global _manager
    if _manager is None:
        return

    await _manager.stop()
    _manager = None


def init_knowledge_manager(
    embedder_model: str,
    embedder_api_key: str | None = None,
    embedder_api_base: str | None = None,
    qdrant_url: str = "http://localhost:16333",
    max_concurrent_jobs: int = 3,
) -> "KnowledgeManager":
    """Create or reuse the singleton knowledge manager."""
    global _manager

    if _manager and _manager.matches_config(
        embedder_model=embedder_model,
        embedder_api_key=embedder_api_key,
        embedder_api_base=embedder_api_base,
        qdrant_url=qdrant_url,
        max_concurrent_jobs=max_concurrent_jobs,
    ):
        _manager.start()
        return _manager

    old_manager = _manager
    if old_manager is not None:
        logger.warning(
            "Reinitializing knowledge manager with a different config; "
            "new uploads will use the new manager instance"
        )
        try:
            asyncio.get_running_loop().create_task(old_manager.stop())
        except RuntimeError:
            pass

    _manager = KnowledgeManager(
        embedder_model=embedder_model,
        embedder_api_key=embedder_api_key,
        embedder_api_base=embedder_api_base,
        qdrant_url=qdrant_url,
        max_concurrent_jobs=max_concurrent_jobs,
    )
    _manager.start()
    return _manager


async def extract_text(data: bytes, filename: str) -> str:
    """Extract text from uploaded bytes. Supports PDF, TXT, MD, and DOCX."""
    name_lower = filename.lower()
    if name_lower.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)

    if name_lower.endswith(".docx"):
        from docx import Document

        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    return data.decode("utf-8", errors="replace")


def _upload_minio(data: bytes, object_name: str) -> None:
    """Upload a file to MinIO. Intended to run inside an executor."""
    from minio import Minio

    from app.config import get_settings

    settings = get_settings()
    client = Minio(
        settings.minio_endpoint,
        access_key=settings.minio_access_key,
        secret_key=settings.minio_secret_key,
        secure=False,
    )

    if not client.bucket_exists(settings.minio_bucket):
        client.make_bucket(settings.minio_bucket)

    client.put_object(settings.minio_bucket, object_name, io.BytesIO(data), len(data))


def _download_minio(object_name: str) -> bytes:
    """Download a file from MinIO. Intended to run inside an executor."""
    from minio import Minio

    from app.config import get_settings

    settings = get_settings()
    client = Minio(
        settings.minio_endpoint,
        access_key=settings.minio_access_key,
        secret_key=settings.minio_secret_key,
        secure=False,
    )

    response = client.get_object(settings.minio_bucket, object_name)
    try:
        return response.read()
    finally:
        response.close()
        response.release_conn()


@dataclass(slots=True)
class KnowledgeTask:
    doc_id: str
    title: str
    content: str


class KnowledgeManager:
    COLLECTION = "ai_os_kb_hybrid_v1"
    DENSE_VECTOR_NAME = "dense"
    SPARSE_VECTOR_NAME = "sparse"
    UPSERT_BATCH = 100
    EMBED_BATCH_SIZE = 5
    EMBED_BATCH_DELAY_SECONDS = 1.0
    RECOVERY_STATUSES = ("pending", "processing")
    TOKEN_PATTERN = re.compile(r"[a-z0-9_]+|[\u4e00-\u9fff]", re.IGNORECASE)

    def __init__(
        self,
        *,
        embedder_model: str,
        embedder_api_key: str | None,
        embedder_api_base: str | None,
        qdrant_url: str,
        max_concurrent_jobs: int,
    ):
        self.embedder_model = embedder_model
        self.embedder_api_key = embedder_api_key
        self.embedder_api_base = embedder_api_base
        self.qdrant_url = qdrant_url
        self.max_concurrent_jobs = max(1, max_concurrent_jobs)
        self.qdrant = AsyncQdrantClient(url=qdrant_url)
        self._vector_size: int | None = None
        self._task_queue: asyncio.Queue[KnowledgeTask] = asyncio.Queue()
        self._worker_tasks: list[asyncio.Task[None]] = []
        self._startup_task: asyncio.Task[None] | None = None
        self._collection_ready = asyncio.Event()
        self._active_jobs = 0
        self._pending_doc_ids: set[str] = set()

    @property
    def queue_size(self) -> int:
        return self._task_queue.qsize()

    @property
    def active_jobs(self) -> int:
        return self._active_jobs

    def matches_config(
        self,
        *,
        embedder_model: str,
        embedder_api_key: str | None,
        embedder_api_base: str | None,
        qdrant_url: str,
        max_concurrent_jobs: int,
    ) -> bool:
        return (
            self.embedder_model == embedder_model
            and self.embedder_api_key == embedder_api_key
            and self.embedder_api_base == embedder_api_base
            and self.qdrant_url == qdrant_url
            and self.max_concurrent_jobs == max(1, max_concurrent_jobs)
        )

    def start(self) -> None:
        if self._worker_tasks:
            return

        self._collection_ready.clear()
        for idx in range(self.max_concurrent_jobs):
            self._worker_tasks.append(
                asyncio.create_task(self._document_worker(idx + 1))
            )
        self._startup_task = asyncio.create_task(self._startup())

    async def stop(self) -> None:
        if self._startup_task is not None:
            self._startup_task.cancel()
            await asyncio.gather(self._startup_task, return_exceptions=True)
            self._startup_task = None

        if self._worker_tasks:
            workers = list(self._worker_tasks)
            self._worker_tasks.clear()
            for worker in workers:
                worker.cancel()
            await asyncio.gather(*workers, return_exceptions=True)

        self._collection_ready.clear()
        self._pending_doc_ids.clear()
        await self.qdrant.close()

    async def store_file_source(self, data: bytes, object_name: str) -> None:
        """Persist an uploaded file before it enters the in-memory queue."""
        loop = asyncio.get_running_loop()
        await loop.run_in_executor(None, _upload_minio, data, object_name)

    async def _load_file_source(self, object_name: str) -> bytes:
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, _download_minio, object_name)

    async def enqueue_document_processing(
        self,
        *,
        doc_id: str,
        title: str,
        content: str,
    ) -> None:
        normalized_content = content.strip()
        if not normalized_content:
            raise ValueError("Document content is empty.")

        if doc_id in self._pending_doc_ids:
            return

        self._pending_doc_ids.add(doc_id)
        await self._task_queue.put(
            KnowledgeTask(
                doc_id=doc_id,
                title=title,
                content=normalized_content,
            )
        )

    async def _startup(self) -> None:
        try:
            collection_created = await self._ensure_collection()
            await self._bootstrap_collection_if_needed(collection_created)
            await self._recover_unfinished_documents()
        except asyncio.CancelledError:
            raise
        except Exception:
            logger.exception("Knowledge manager startup failed")
        finally:
            self._collection_ready.set()

    async def _document_worker(self, worker_id: int) -> None:
        while True:
            task = await self._task_queue.get()
            self._active_jobs += 1
            try:
                await self._collection_ready.wait()
                logger.info(
                    "Knowledge worker %s picked doc %s (queued=%s)",
                    worker_id,
                    task.doc_id,
                    self._task_queue.qsize(),
                )
                await self._process_task(task)
            except asyncio.CancelledError:
                raise
            except Exception:
                logger.exception("Unhandled knowledge worker error for doc %s", task.doc_id)
                await self._mark_document_error(task.doc_id, "Unhandled background processing error")
            finally:
                self._active_jobs = max(0, self._active_jobs - 1)
                self._pending_doc_ids.discard(task.doc_id)
                self._task_queue.task_done()

    async def _process_task(self, task: KnowledgeTask) -> None:
        await self.process_document(task.doc_id, task.content, title=task.title)

    async def _mark_document_error(self, doc_id: str, error_msg: str) -> None:
        from app.core.database import AsyncSessionLocal
        from app.models.knowledge import KnowledgeDocument

        async with AsyncSessionLocal() as db:
            await db.execute(
                update(KnowledgeDocument)
                .where(KnowledgeDocument.id == doc_id)
                .values(status="error", error_msg=error_msg[:500])
            )
            await db.commit()

    def _chunk_text(self, text: str, size: int = 500, overlap: int = 100) -> list[str]:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
        if not normalized:
            return []

        if len(normalized) <= size:
            return [normalized]

        chunks: list[str] = []
        start = 0

        while start < len(normalized):
            end = min(len(normalized), start + size)
            if end < len(normalized):
                newline_break = normalized.rfind("\n", start, end)
                space_break = normalized.rfind(" ", start, end)
                boundary = max(newline_break, space_break)
                if boundary > start + size // 2:
                    end = boundary

            chunk = normalized[start:end].strip()
            if chunk:
                chunks.append(chunk)

            if end >= len(normalized):
                break

            next_start = end - overlap if overlap > 0 else end
            if next_start <= start:
                next_start = end
            start = next_start

        return chunks

    def _tokenize_for_sparse(self, text: str) -> list[str]:
        return self.TOKEN_PATTERN.findall(text.lower())

    def _term_to_sparse_index(self, term: str) -> int:
        digest = hashlib.blake2b(term.encode("utf-8"), digest_size=8).digest()
        return int.from_bytes(digest, "big") % 2_147_483_647

    def _build_sparse_vector(self, text: str) -> SparseVector | None:
        tokens = self._tokenize_for_sparse(text)
        if not tokens:
            return None

        weighted_terms: dict[int, float] = {}
        for term, count in Counter(tokens).items():
            sparse_index = self._term_to_sparse_index(term)
            weighted_terms[sparse_index] = weighted_terms.get(sparse_index, 0.0) + (1.0 + math.log(count))

        items = sorted(weighted_terms.items())
        if not items:
            return None

        return SparseVector(
            indices=[index for index, _ in items],
            values=[value for _, value in items],
        )

    async def _embed_one(
        self,
        text: str,
        client: "httpx.AsyncClient",
        url: str,
        headers: dict[str, str],
    ) -> list[float]:
        import httpx as _httpx

        payload = {"model": self.embedder_model, "input": text}
        last_error = ""

        for attempt in range(5):
            try:
                response = await client.post(url, json=payload, headers=headers)
            except (_httpx.ReadTimeout, _httpx.ConnectTimeout, _httpx.ConnectError) as exc:
                wait_seconds = 2**attempt
                last_error = str(exc)
                logger.warning(
                    "Knowledge embedding network error (%s), retrying in %ss",
                    type(exc).__name__,
                    wait_seconds,
                )
                await asyncio.sleep(wait_seconds)
                continue

            if response.status_code == 429:
                wait_seconds = 2**attempt
                last_error = response.text
                logger.warning(
                    "Knowledge embedding rate limited, retrying in %ss (attempt %s/5)",
                    wait_seconds,
                    attempt + 1,
                )
                await asyncio.sleep(wait_seconds)
                continue

            if response.status_code >= 400:
                logger.error(
                    "Knowledge embedding API error %s: %s",
                    response.status_code,
                    response.text,
                )
                response.raise_for_status()

            return response.json()["data"][0]["embedding"]

        raise RuntimeError(f"Failed after 5 retries: {last_error}")

    async def _embed_batch(self, texts: list[str]) -> list[list[float]]:
        if self.embedder_api_base:
            import httpx

            base = self.embedder_api_base.rstrip("/")
            url = f"{base}/embeddings"
            headers = {"Content-Type": "application/json"}
            if self.embedder_api_key:
                headers["Authorization"] = f"Bearer {self.embedder_api_key}"

            logger.info("Embedding %s knowledge chunks", len(texts))
            results: list[list[float]] = []

            async with httpx.AsyncClient(timeout=60) as client:
                for start in range(0, len(texts), self.EMBED_BATCH_SIZE):
                    batch = texts[start:start + self.EMBED_BATCH_SIZE]
                    batch_results = await asyncio.gather(
                        *[self._embed_one(text, client, url, headers) for text in batch]
                    )
                    results.extend(batch_results)

                    if start + self.EMBED_BATCH_SIZE < len(texts):
                        await asyncio.sleep(self.EMBED_BATCH_DELAY_SECONDS)

                    logger.info(
                        "Knowledge embedding progress %s/%s",
                        min(start + self.EMBED_BATCH_SIZE, len(texts)),
                        len(texts),
                    )

            return results

        from litellm import aembedding

        response = await aembedding(
            model=self.embedder_model,
            input=texts,
            api_key=self.embedder_api_key,
        )
        return [item["embedding"] for item in response.data]

    async def _ensure_collection(self) -> bool:
        if await self.qdrant.collection_exists(self.COLLECTION):
            info = await self.qdrant.get_collection(self.COLLECTION)
            vectors = info.config.params.vectors
            sparse_vectors = info.config.params.sparse_vectors or {}

            if isinstance(vectors, dict) and self.DENSE_VECTOR_NAME in vectors and self.SPARSE_VECTOR_NAME in sparse_vectors:
                self._vector_size = vectors[self.DENSE_VECTOR_NAME].size
                return False

            logger.warning(
                "Collection %s has incompatible vector config; recreating hybrid index",
                self.COLLECTION,
            )
            await self.qdrant.delete_collection(self.COLLECTION)

        if self._vector_size is None:
            probe_embedding = (await self._embed_batch(["knowledge bootstrap probe"]))[0]
            self._vector_size = len(probe_embedding)

        await self.qdrant.create_collection(
            collection_name=self.COLLECTION,
            vectors_config={
                self.DENSE_VECTOR_NAME: VectorParams(
                    size=self._vector_size,
                    distance=Distance.COSINE,
                )
            },
            sparse_vectors_config={
                self.SPARSE_VECTOR_NAME: SparseVectorParams(modifier=Modifier.IDF)
            },
        )
        logger.info(
            "Created Qdrant hybrid collection %s (dense_dim=%s)",
            self.COLLECTION,
            self._vector_size,
        )
        return True

    async def create_document(
        self,
        *,
        title: str,
        content: str,
        source_type: str,
        db: "AsyncSession",
        source_url: str | None = None,
    ):
        from datetime import datetime, timezone

        from app.models.knowledge import KnowledgeDocument

        normalized_content = content.strip()
        if not normalized_content:
            raise ValueError("Document content is empty.")

        doc_id = str(uuid.uuid4())
        doc = KnowledgeDocument(
            id=doc_id,
            title=title,
            source_type=source_type,
            source_url=source_url,
            raw_content=normalized_content,
            chunk_count=0,
            status="pending",
            created_at=datetime.now(timezone.utc),
        )
        db.add(doc)
        await db.flush()
        await db.commit()
        return doc_id, doc

    async def _delete_points_for_doc(self, doc_id: str) -> None:
        if not await self.qdrant.collection_exists(self.COLLECTION):
            return

        await self.qdrant.delete(
            collection_name=self.COLLECTION,
            points_selector=Filter(
                must=[
                    FieldCondition(
                        key="doc_id",
                        match=MatchValue(value=doc_id),
                    )
                ]
            ),
        )

    async def _clear_document_artifacts(self, doc_id: str, db: "AsyncSession") -> None:
        from app.models.knowledge import KnowledgeChunk

        await self._delete_points_for_doc(doc_id)
        await db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.doc_id == doc_id))

    async def _index_chunks(
        self,
        *,
        db: "AsyncSession",
        doc_id: str,
        title: str,
        chunks: list[str],
        existing_chunk_rows: list | None = None,
    ) -> int:
        from app.models.knowledge import KnowledgeChunk

        if not chunks:
            raise ValueError("Document did not produce any chunks.")

        embeddings = await self._embed_batch(chunks)
        if not embeddings:
            raise ValueError("Embedding provider returned no vectors.")

        if self._vector_size is None:
            self._vector_size = len(embeddings[0])
        await self._ensure_collection()

        existing_by_index = {
            chunk_row.chunk_index: chunk_row for chunk_row in (existing_chunk_rows or [])
        }
        reused_chunk_ids: set[str] = set()
        new_chunk_models: list[KnowledgeChunk] = []
        points: list[PointStruct] = []

        for index, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
            point_id = str(uuid.uuid4())
            sparse_vector = self._build_sparse_vector(chunk_text)
            point_vector: dict[str, list[float] | SparseVector] = {
                self.DENSE_VECTOR_NAME: embedding,
            }
            if sparse_vector is not None:
                point_vector[self.SPARSE_VECTOR_NAME] = sparse_vector

            points.append(
                PointStruct(
                    id=point_id,
                    vector=point_vector,
                    payload={
                        "doc_id": doc_id,
                        "title": title,
                        "chunk_index": index,
                        "content": chunk_text,
                    },
                )
            )

            existing_chunk = existing_by_index.get(index)
            if existing_chunk is None:
                new_chunk_models.append(
                    KnowledgeChunk(
                        id=str(uuid.uuid4()),
                        doc_id=doc_id,
                        content=chunk_text,
                        chunk_index=index,
                        qdrant_point_id=point_id,
                    )
                )
                continue

            existing_chunk.content = chunk_text
            existing_chunk.chunk_index = index
            existing_chunk.qdrant_point_id = point_id
            reused_chunk_ids.add(existing_chunk.id)

        stale_chunk_ids = [
            chunk_row.id
            for chunk_row in (existing_chunk_rows or [])
            if chunk_row.id not in reused_chunk_ids
        ]
        if stale_chunk_ids:
            await db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.id.in_(stale_chunk_ids)))

        for start in range(0, len(points), self.UPSERT_BATCH):
            batch = points[start:start + self.UPSERT_BATCH]
            await self.qdrant.upsert(collection_name=self.COLLECTION, points=batch)
            logger.info(
                "Knowledge upsert progress %s/%s for doc %s",
                min(start + self.UPSERT_BATCH, len(points)),
                len(points),
                doc_id,
            )

        for chunk in new_chunk_models:
            db.add(chunk)

        return len(chunks)

    async def _reindex_document_from_chunk_rows(
        self,
        doc_id: str,
        title: str,
    ) -> None:
        from app.core.database import AsyncSessionLocal
        from app.models.knowledge import KnowledgeChunk, KnowledgeDocument

        async with AsyncSessionLocal() as db:
            document = await db.get(KnowledgeDocument, doc_id)
            if document is None:
                return

            result = await db.execute(
                select(KnowledgeChunk)
                .where(KnowledgeChunk.doc_id == doc_id)
                .order_by(KnowledgeChunk.chunk_index.asc())
            )
            chunk_rows = result.scalars().all()
            if not chunk_rows:
                return

            chunk_count = await self._index_chunks(
                db=db,
                doc_id=doc_id,
                title=title,
                chunks=[chunk.content for chunk in chunk_rows],
                existing_chunk_rows=chunk_rows,
            )
            await db.execute(
                update(KnowledgeDocument)
                .where(KnowledgeDocument.id == doc_id)
                .values(status="done", chunk_count=chunk_count, error_msg=None)
            )
            await db.commit()

    async def _bootstrap_collection_if_needed(self, collection_created: bool) -> None:
        if not collection_created:
            info = await self.qdrant.get_collection(self.COLLECTION)
            if (info.points_count or 0) > 0:
                return

        from app.core.database import AsyncSessionLocal
        from app.models.knowledge import KnowledgeDocument

        async with AsyncSessionLocal() as db:
            result = await db.execute(
                select(KnowledgeDocument)
                .where(KnowledgeDocument.status == "done")
                .order_by(KnowledgeDocument.created_at.asc())
            )
            completed_documents = result.scalars().all()

        if not completed_documents:
            return

        logger.info(
            "Bootstrapping hybrid collection from %s completed documents",
            len(completed_documents),
        )
        for document in completed_documents:
            if document.raw_content and document.raw_content.strip():
                await self.process_document(document.id, document.raw_content, title=document.title)
                continue

            await self._reindex_document_from_chunk_rows(document.id, document.title)

    async def _recover_unfinished_documents(self) -> None:
        from app.core.database import AsyncSessionLocal
        from app.models.knowledge import KnowledgeDocument

        async with AsyncSessionLocal() as db:
            result = await db.execute(
                select(KnowledgeDocument)
                .where(KnowledgeDocument.status.in_(self.RECOVERY_STATUSES))
                .order_by(KnowledgeDocument.created_at.asc())
            )
            documents = result.scalars().all()

        if not documents:
            return

        logger.info("Recovering %s unfinished knowledge documents", len(documents))
        for document in documents:
            content = await self._recover_document_content(document)
            if not content:
                await self._mark_document_error(
                    document.id,
                    "Unable to recover source content after restart.",
                )
                continue

            await self.enqueue_document_processing(
                doc_id=document.id,
                title=document.title,
                content=content,
            )

    async def _recover_document_content(self, document) -> str | None:
        if document.raw_content and document.raw_content.strip():
            return document.raw_content

        if document.source_type != "file" or not document.source_url:
            return None

        try:
            file_bytes = await self._load_file_source(document.source_url)
            content = await extract_text(file_bytes, document.title)
        except Exception:
            logger.exception("Failed to reload source file for knowledge doc %s", document.id)
            return None

        normalized_content = content.strip()
        if not normalized_content:
            return None

        from app.core.database import AsyncSessionLocal
        from app.models.knowledge import KnowledgeDocument

        async with AsyncSessionLocal() as db:
            await db.execute(
                update(KnowledgeDocument)
                .where(KnowledgeDocument.id == document.id)
                .values(raw_content=normalized_content)
            )
            await db.commit()

        return normalized_content

    async def process_document(self, doc_id: str, content: str, title: str = "") -> None:
        """Process one document: chunk, embed, store in Qdrant, update DB."""
        from app.core.database import AsyncSessionLocal
        from app.models.knowledge import KnowledgeDocument

        normalized_content = content.strip()
        if not normalized_content:
            await self._mark_document_error(doc_id, "Document content is empty.")
            return

        async with AsyncSessionLocal() as db:
            try:
                existing = await db.get(KnowledgeDocument, doc_id)
                if existing is None:
                    logger.info("Skip processing deleted knowledge doc %s", doc_id)
                    return

                await db.execute(
                    update(KnowledgeDocument)
                    .where(KnowledgeDocument.id == doc_id)
                    .values(
                        status="processing",
                        error_msg=None,
                        raw_content=normalized_content,
                    )
                )
                await db.commit()

                chunks = self._chunk_text(normalized_content)
                logger.info("Processing knowledge doc %s with %s chunks", doc_id, len(chunks))

                await self._clear_document_artifacts(doc_id, db)
                chunk_count = await self._index_chunks(
                    db=db,
                    doc_id=doc_id,
                    title=title,
                    chunks=chunks,
                )

                await db.execute(
                    update(KnowledgeDocument)
                    .where(KnowledgeDocument.id == doc_id)
                    .values(status="done", chunk_count=chunk_count, error_msg=None)
                )
                await db.commit()
                logger.info("Finished knowledge doc %s", doc_id)

            except Exception as exc:
                logger.exception("Error processing knowledge doc %s", doc_id)
                await db.rollback()
                await db.execute(
                    update(KnowledgeDocument)
                    .where(KnowledgeDocument.id == doc_id)
                    .values(status="error", error_msg=(str(exc) or type(exc).__name__)[:500])
                )
                await db.commit()

    async def search(self, query: str, limit: int = 5) -> list[dict]:
        query_text = query.strip()
        if not query_text:
            return []

        if not await self.qdrant.collection_exists(self.COLLECTION):
            return []

        dense_embedding = (await self._embed_batch([query_text]))[0]
        sparse_query = self._build_sparse_vector(query_text)
        prefetch_limit = max(limit * 4, limit)

        if sparse_query is None:
            response = await self.qdrant.query_points(
                collection_name=self.COLLECTION,
                query=dense_embedding,
                using=self.DENSE_VECTOR_NAME,
                limit=limit,
                with_payload=True,
                score_threshold=0.3,
            )
        else:
            response = await self.qdrant.query_points(
                collection_name=self.COLLECTION,
                prefetch=[
                    Prefetch(
                        query=dense_embedding,
                        using=self.DENSE_VECTOR_NAME,
                        limit=prefetch_limit,
                    ),
                    Prefetch(
                        query=sparse_query,
                        using=self.SPARSE_VECTOR_NAME,
                        limit=prefetch_limit,
                    ),
                ],
                query=FusionQuery(fusion=Fusion.RRF),
                limit=limit,
                with_payload=True,
            )

        return [
            {
                "content": point.payload.get("content", ""),
                "title": point.payload.get("title", ""),
                "doc_id": point.payload.get("doc_id", ""),
                "chunk_index": point.payload.get("chunk_index", 0),
                "score": point.score,
            }
            for point in response.points
        ]

    async def delete_document(self, doc_id: str, db: "AsyncSession") -> None:
        from app.models.knowledge import KnowledgeChunk, KnowledgeDocument

        self._pending_doc_ids.discard(doc_id)
        await self._delete_points_for_doc(doc_id)
        await db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.doc_id == doc_id))
        await db.execute(delete(KnowledgeDocument).where(KnowledgeDocument.id == doc_id))
        logger.info("Deleted knowledge doc %s", doc_id)
